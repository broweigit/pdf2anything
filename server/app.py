from flask import Flask, request, current_app, jsonify, make_response, session, send_from_directory, send_file, render_template
from flask_cors import CORS
from flask_session import Session
from flask_sqlalchemy import SQLAlchemy

from ocr import OCRSystem
from chat import ChatSystem
from fileman import FileManSystem, make_base64_png, convert_image_to_binary, convert_pdf_to_binary, restore_binary_to_image, restore_binary_to_pdf

import cv2
import numpy as np
import base64
import click

import sys
import os
import random
import io
import shutil

from docx import Document
from docxcompose.composer import Composer

# from ruler import calculate_line_length

def create_app():
    # SQLite 配置
    WIN = sys.platform.startswith('win')
    if WIN:  # 如果是 Windows 系统，使用三个斜线
        prefix = 'sqlite:///'
    else:  # 否则使用四个斜线
        prefix = 'sqlite:////'

    app = Flask(__name__)
    app.config['SECRET_KEY'] = '12010524'
    app.config['SESSION_TYPE'] = 'filesystem'  # 设置 session 存储方式为文件系统(也可以利用sqlalchemy)
    app.config['SESSION_PERMANENT'] = True  # 设置会话为永久性
    app.config['SESSION_USE_SIGNER'] = False  # 是否对发送到浏览器上session的cookie值进行加密
    app.config['SESSION_KEY_PREFIX'] = 'session:'  # 保存到session中的值的前缀
    app.config['SQLALCHEMY_DATABASE_URI'] = prefix + os.path.join(app.root_path, 'data.db')
    app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False  # 关闭对模型修改的监控
    db = SQLAlchemy(app)
    Session(app)
    CORS(app, supports_credentials=True)

    ocr = OCRSystem()
    
    with app.app_context():
        current_app.chat = ChatSystem()
        current_app.fm = FileManSystem()

    # Table
    class User(db.Model):  # 表名将会是 user（自动生成，小写处理）
        id = db.Column(db.Integer, primary_key=True)  # 主键 id
        username = db.Column(db.String(20))  # 名字
        account = db.Column(db.String(9))  # 账号 固定9位
        password = db.Column(db.String(30))  # 密码

    class Project(db.Model):
        id = db.Column(db.Integer, primary_key=True)
        userId = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
        name = db.Column(db.String(100), nullable=False)
        jsonData = db.Column(db.Text, nullable=False)
        projectDate = db.Column(db.String(20), nullable=False)
        file_id = db.Column(db.Integer, db.ForeignKey('file.id'), nullable=False)

        # Define the relationship with User model(one to many)
        user = db.relationship('User', backref=db.backref('projects', lazy=True))

        # Define the relationship with File model(one to one)
        file = db.relationship('File', backref=db.backref('project', lazy=True))

    class File(db.Model):
        id = db.Column(db.Integer, primary_key=True)
        pdfData = db.Column(db.BLOB)
        imgListData = db.Column(db.PickleType)

    @app.cli.command()  # 注册为命令，可以传入 name 参数来自定义命令
    @click.option('--drop', is_flag=True, help='Create after drop.')  # 设置选项
    def initdb(drop):
        """Initialize the database."""
        if drop:  # 判断是否输入了选项
            db.drop_all()
        db.create_all()
        click.echo('Initialized database.')  # 输出提示信息

    @app.route('/')
    def index():
        return render_template('index.html')  # 渲染打包好的React App的页面
    
    @app.route('/<path:filename>')
    def manifest(filename):
        return send_from_directory('./templates', filename)
    
    # @app.route('/<path:filename>')
    # def root(filename):
    #     return send_from_directory('./public', filename)

    @app.route('/upload-image', methods=['POST'])
    def upload_image():
        with app.app_context():

            img_file = request.files['image']

            return current_app.fm.loadImg(img_file)
        
    @app.route('/upload-pdf', methods=['POST'])
    def upload_pdf():
        with app.app_context():

            file = request.files['file']  # 获取上传的文件对象
            if file and file.filename.endswith('.pdf'):
                return current_app.fm.loadPdf(file)
            
    @app.route('/upload-reset', methods=['POST'])
    def upload_reset():
        with app.app_context():
            return current_app.fm.reset()

    @app.route('/layout-analysis', methods=['GET'])
    def get_layout_analysis():
        with app.app_context():
            page_id = request.args.get('pageId') 
            layout_result = ocr.layout_analysis(current_app.fm.img_list[int(page_id)])

            return jsonify(layout_result)
    
    @app.route('/table-extract', methods=['GET'])
    def get_table_extract():
        with app.app_context():
            bbox = request.args.get('bbox')
            page_id = request.args.get('pageId') 
            bbox = [int(float(i)) for i in bbox[1:-1].split(',')]

            x, y, w, h = bbox
            img_height, img_width = current_app.fm.img_list[int(page_id)].shape[:2]

            # Adjust x, y, w, h if they exceed image boundaries
            x = max(x, 0)
            y = max(y, 0)
            w = min(w, img_width - x)
            h = min(h, img_height - y)
            img_roi = current_app.fm.img_list[int(page_id)][y:y+h, x:x+w]

            # cv2.imwrite('./roi.jpg', img_roi)

            table_result = ocr.table_analysis(img_roi)
            # print(table_result[0]['res']['html'])
            response = make_response(table_result[0]['res']['html'])
            response.headers['Content-Type'] = 'text/html'
            return response
    
    @app.route('/text-extract', methods=['GET'])
    def get_text_extract():
        with app.app_context():
            bbox = request.args.get('bbox') 
            lang = request.args.get('lang') 
            page_id = request.args.get('pageId') 
            bbox = [int(float(i)) for i in bbox[1:-1].split(',')] 

            x, y, w, h = bbox
            img_height, img_width = current_app.fm.img_list[int(page_id)].shape[:2]

            # Adjust x, y, w, h if they exceed image boundaries
            x = max(x, 0)
            y = max(y, 0)
            w = min(w, img_width - x)
            h = min(h, img_height - y)
            img_roi = current_app.fm.img_list[int(page_id)][y:y+h, x:x+w]

            # cv2.imwrite('./roi.jpg', img_roi)

            text_result = ocr.text_analysis(img_roi, lang)
            return text_result
    
    @app.route('/convert-image', methods=['POST'])
    def convert_image():
        # Get the image file from the request
        image_file = request.files['image']
        
        # Read the image with OpenCV
        image = cv2.imdecode(np.fromstring(image_file.read(), np.uint8), cv2.IMREAD_COLOR)
        
        # Convert the image to PNG format
        _, png_data = cv2.imencode('.png', image)
        
        # Encode the PNG data as base64
        base64_data = base64.b64encode(png_data)
        base64_data_str = "data:image/png;base64," + base64_data.decode("utf-8")
        
        # Return the base64 string as the response
        return base64_data_str
    
    
    @app.route('/chat-stream', methods=['GET'])
    def chat_stream():
        with app.app_context():
            prompt = request.args.get('prompt') 
            return current_app.chat.stream_response(prompt)
    
    @app.route('/chat-reset', methods=['POST'])
    def chat_reset():
        with app.app_context():
            return current_app.chat.reset()
        
    # 登录
    @app.route('/user/login', methods=['GET'])
    def login():
        # 获取表单数据
        account = request.args.get('account')  # 传入表单对应输入字段的 name 值
        password = request.args.get('password')  # 传入表单对应输入字段的 password 值

        user = User.query.filter_by(account=account, password=password).first()
        if user is None:
            # 尝试以用户名登录
            user = User.query.filter_by(username=account, password=password).first()
            if user is None: 
                return make_response('Login Failed', 400)
            
        # session['user_id'] = user.id

        normal_response = jsonify({'userId': user.id, 'username': user.username, 'account': user.account})
        return make_response(normal_response, 200)

    # 注册
    @app.route('/user/register', methods=['GET'])
    def register():
        # 获取表单数据
        username = request.args.get('username')  # 传入表单对应输入字段的 name 值
        password = request.args.get('password')  # 传入表单对应输入字段的 password 值

        print(username, password)

        # 生成账号
        num = 0
        while True:
            num = random.randint(100000000, 999999999)  # 闭区间
            if User.query.filter_by(account=str(num)).first() is None:
                break
        user = User(username=username, account=str(num), password=password)  # 创建记录
        db.session.add(user)  # 添加到数据库会话
        db.session.commit()  # 提交数据库会话
        normal_response = jsonify({'account': str(num)})
        return make_response(normal_response, 200)
    
    # 退出登录
    @app.route('/user/logout', methods=['GET'])
    def logout():
        # print(session.get('user_id'), session.sid)
        # session.pop('user_id', None)
        return make_response('已退出登录', 200)
    
    @app.route('/save-project', methods=['POST'])
    def save_project():
        # if 'user_id' not in session or session['user_id'] is None:
        #     return make_response('User not logged in', 401)
        
        project_name = request.form.get('projectName')  # 获取表单字段 projectName 的值
        json_data = request.form.get('jsonData')  # 获取表单字段 jsonData 的值
        project_date = request.form.get('projectDate')  # 获取表单字段 projectDate 的值

        pdf_file = current_app.fm.pdf_file
        img_list = current_app.fm.img_list

        # user_id = session['user_id']
        user_id = request.form.get('userId')

        # 保存 pdf_file和img_list 到 File 表
        pdf_file_data = convert_pdf_to_binary(pdf_file)
        img_list_data = [convert_image_to_binary(image) for image in img_list]
        file = File(pdfData=pdf_file_data, imgListData=img_list_data)
        db.session.add(file)
        db.session.commit()

        # 获取文件 ID
        file_id = file.id

        # 创建 Project 记录并关联文件
        project = Project(userId=user_id, name=project_name, jsonData=json_data, projectDate=project_date, file_id=file_id)
        db.session.add(project)
        db.session.commit()

        # 构造返回结果
        response_data = {
            'saved': True
        }
        response = jsonify(response_data)
        return make_response(response, 200)
    
    @app.route('/get-project-list', methods=['GET'])
    def get_project_list():
        user_id = request.args.get('userId')  # 获取用户ID参数

        # 查询特定用户的所有项目
        projects = Project.query.filter_by(userId=user_id).all()

        # 构造返回的项目列表
        project_list = []
        for project in projects:
            project_data = {
                'pid': project.id,
                'name': project.name,
                'projectDate': project.projectDate
            }
            project_list.append(project_data)

        return jsonify(project_list)
    
    @app.route('/load-project', methods=['GET'])
    def load_project():
        project_id = request.args.get('pid')  # 获取表单字段 projectId 的值

        # 查询项目记录
        project = Project.query.get(project_id)

        if not project:
            return make_response('Project not found', 404)

        # 查询关联的文件记录
        file = File.query.get(project.file_id)

        if not file:
            return make_response('File not found', 404)

        # 将文件数据转换为相应格式
        pdf = restore_binary_to_pdf(file.pdfData)
        img_list = [restore_binary_to_image(image_data) for image_data in file.imgListData]
        # 保存到filemanager中
        current_app.fm.pdf_file = pdf
        current_app.fm.img_list = img_list

        # 构造返回结果
        response_data = {
            'jsonData': project.jsonData,
            'imgList': make_base64_png(img_list)
        }
        response = jsonify(response_data)
        return make_response(response, 200)
    
    @app.route('/delete-project', methods=['GET'])
    def delete_project():
        project_id = request.args.get('pid')  # 获取表单字段 projectId 的值

        project = Project.query.get(project_id)

        if project:
            # 删除项目记录
            db.session.delete(project) 
            db.session.commit()
            return make_response('删除成功', 200)
        else:
            return make_response('项目不存在', 404)
        
    @app.route('/convert-file', methods=['GET'])
    def convert_file():
        file_type = request.args.get('fileType')

        if file_type == 'png':
            return make_response(make_base64_png(current_app.fm.img_list), 200)
        elif file_type == 'pdf':
            # 先将img列表转换
            current_app.fm.imgs2pdf()
            pdf_data = current_app.fm.pdf_file.tobytes()

            # 将PDF文件进行Base64编码
            encoded_pdf = base64.b64encode(pdf_data).decode('utf-8')

            # 返回Base64编码后的PDF数据给前端
            return jsonify({'pdfData': encoded_pdf})
        elif file_type == 'doc':
            index = 0
            random_number = random.randint(10000000, 99999999)
            folder_name = './doc_recovery/'+ str(random_number)
            master = folder_name + '/temp' + '0' + '_ocr.docx'
            for img in current_app.fm.img_list:
                ocr.recovery(img, folder_name, index)
                if index > 0:
                    sub = folder_name + '/temp' + str(index) + '_ocr.docx'
                    doc_master = Document(master)
                    doc_master.add_page_break()
                    cp = Composer(doc_master)
                    cp.append(Document(sub))
                    doc_master.save(master)
                index = index + 1

            doc_master = Document(master)
            # 创建一个字节流对象
            output = io.BytesIO()
            # 将doc_master保存到字节流中
            doc_master.save(output)
            # 将字节流转换为字节数据
            doc_data = output.getvalue()
            # 将doc文件进行Base64编码
            encoded_doc = base64.b64encode(doc_data).decode('utf-8')

            # 使用shutil.rmtree函数删除文件夹及其内容
            shutil.rmtree(folder_name)

            # 返回Base64编码后的doc数据给前端
            return jsonify({'docData': encoded_doc})

        else:
            return make_response('不支持此格式转换', 404)

    return app

if __name__ == '__main__':
    app = create_app()
    app.run(debug=False, host='0.0.0.0', port='5000')